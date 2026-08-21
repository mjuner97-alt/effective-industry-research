#!/usr/bin/env python3
"""Part 1: Create base document with cover page + chapters 1-2"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTDIR = os.path.dirname(os.path.abspath(__file__))
doc = Document()

# ---- Styles ----
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level, size, bold in [(1,18,True),(2,16,True),(3,14,True)]:
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.font.size = Pt(size)
    h.font.bold = bold
    h.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return table

# ---- Cover Page ----
for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('缺陷分类标签体系 + 缺陷归因分析')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('业界技术调研报告')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x44,0x44,0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Industry Research Report')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x88,0x88,0x99)

for _ in range(3):
    doc.add_paragraph()

info_lines = [
    '调研日期：2026年7月',
    '调研方法：五维调研法（学术界+大厂+企业+银行+大会）',
    '来源数量：共64个有效来源',
    '根据 effective-industry-research skill 方法论生成',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(12)

doc.add_page_break()

# ---- Chapter 1 ----
doc.add_heading('第一章 调研背景与目标', level=1)

doc.add_heading('1.1 调研主题', level=2)
doc.add_paragraph('缺陷分类标签体系 + 缺陷问题归因分析（Defect Classification & Labeling System + Defect Root Cause Analysis）')

doc.add_heading('1.2 调研背景与动机', level=2)
doc.add_paragraph('在软件研发与测试效能领域，缺陷管理是质量保障的核心环节。然而许多团队面临以下痛点：')
for pain in [
    '缺陷分类混乱：不同团队使用不同分类标准，无法横向比较和汇聚分析',
    '归因分析缺失：缺陷修复后缺少根因分析，同类问题反复出现',
    '度量体系不全：缺乏统一的缺陷标签体系，难以量化评估质量改进效果',
    '改进闭环断裂：从缺陷发现到流程改进缺少数据驱动的闭环机制',
]:
    doc.add_paragraph(pain, style='List Bullet')

doc.add_paragraph('本次调研要回答的核心问题：')
for q in [
    '业界主流的缺陷分类体系有哪些？如何选择适合团队的分类标准？',
    '大厂和银行在缺陷归因分析上有哪些成熟实践？',
    '如何建立从缺陷分类到归因分析再到流程改进的闭环？',
]:
    doc.add_paragraph(q, style='List Number')

doc.add_heading('1.3 调研范围', level=2)
doc.add_paragraph('时间范围：近5年（2020-2026），经典论文追溯至1992年')
doc.add_paragraph('覆盖维度：学术界文献、互联网大厂实践、国外企业方案、银行同业案例、业界大会')
doc.add_paragraph('边界说明：本调研聚焦软件缺陷的分类与归因，不涉及硬件缺陷、安全漏洞专项分类（CWE/CVE体系仅作参考）')

doc.add_heading('1.4 调研方法论', level=2)
doc.add_paragraph('采用五维调研法，对主题从五个维度进行系统化调研：')
add_table(doc,
    ['维度', '调研目标', '核心渠道'],
    [
        ['学术界文献', '理论框架、标准体系、分类法', 'Google Scholar, IEEE, 知网'],
        ['互联网大厂', '业界领先公司的实际做法', '美团/阿里/百度/京东/掘金等技术博客'],
        ['国外企业方案', '商业产品、开源项目、行业报告', 'Jira/Sentry/DORA/Google SRE/GitLab'],
        ['银行同业案例', '金融行业落地、监管要求', '银保监会/银行技术论文/金融科技研究'],
        ['业界大会', '最前沿趋势和实践', 'QCon/SREcon/ICSE/QECon等'],
    ],
    col_widths=[3, 5, 5]
)

doc.add_heading('1.5 关键词矩阵', level=2)
add_table(doc,
    ['维度', '关键词组', '搜索内容说明'],
    [
        ['学术界', 'ODC正交缺陷分类 / IEEE 1044 / bug taxonomy survey / RCA root cause analysis / defect prediction ML', '经典框架+最新研究+综述'],
        ['大厂', '美团+缺陷管理+质量 / 阿里+蚂蚁+缺陷归因 / 字节+华为+质量保障 / Google SRE postmortem', '站内搜索+技术博客'],
        ['国外企业', 'Jira defect workflow / Sentry issue grouping / DORA report / ServiceNow problem management', '产品文档+行业报告'],
        ['银行同业', '银保监会信息科技指引 / 银行+质量保障+DevOps / 金融科技风险管理', '监管文件+银行技术论文'],
        ['大会', 'QECon+缺陷管理 / QCon+质量 / SREcon+RCA / ICSE+defect prediction', '大会议程+演讲材料'],
    ],
    col_widths=[2.5, 5.5, 5]
)

doc.add_page_break()

# ---- Chapter 2 ----
doc.add_heading('第二章 学术界文献综述', level=1)

doc.add_heading('2.1 经典理论与框架', level=2)

doc.add_heading('2.1.1 ODC正交缺陷分类法（Orthogonal Defect Classification）', level=3)
doc.add_paragraph('ODC是由IBM Ram Chillarege于1992年提出的缺陷分类体系，是缺陷分类领域最经典、影响最深远的框架之一。[1]')
p = doc.add_paragraph()
p.add_run('提出者：').bold = True
p.add_run('Ram Chillarege (IBM)')
p = doc.add_paragraph()
p.add_run('提出时间：').bold = True
p.add_run('1992年')
p = doc.add_paragraph()
p.add_run('核心思想：').bold = True
p.add_run('将缺陷按8个正交维度进行分类，每个维度内部类别互不重叠（正交），从而实现对缺陷的多维度定量分析，支持精确的过程改进指导。')
p = doc.add_paragraph()
p.add_run('分类维度（8维度114类）：').bold = True
for dim in [
    'Activity（发现活动）：8类 — 如单元测试、功能测试、系统测试等',
    'Trigger（触发因素）：36类 — 如边界条件、变异覆盖等',
    'Impact（结果影响）：13类 — 如容错能力、可安装性等',
    'Target（问题根源对象）：6类 — 如设计/代码/文档等',
    'Type（缺陷类型）：39类 — 如赋值错误/接口错误等',
    'Qualifier（缺陷定界）：3类 — 如缺失/错误/额外',
    'Source（责任来源）：5类 — 如内部开发/外部供应商等',
    'Age（缺陷年龄）：4类 — 如新开发/基线等',
]:
    doc.add_paragraph(dim, style='List Bullet')
p = doc.add_paragraph()
p.add_run('优点：').bold = True
p.add_run('① 多维度正交分类支持定量分析；② 与过程改进直接关联；③ 业界验证充分')
p = doc.add_paragraph()
p.add_run('缺点：').bold = True
p.add_run('① 分类维度较多，实施门槛高；② 需要专门培训分类人员；③ 不同项目间分类标准需调整')
p = doc.add_paragraph()
p.add_run('适用场景：').bold = True
p.add_run('中大型软件组织的过程改进与质量度量')
p = doc.add_paragraph()
p.add_run('成熟度：🟢 已规模化').bold = True

doc.add_heading('2.1.2 根因分析法（Root Cause Analysis, RCA）', level=3)
doc.add_paragraph('RCA是一种结构化的问题根因识别方法，起源于制造业，后广泛应用于软件缺陷归因分析。[2][59]')

p = doc.add_paragraph()
p.add_run('核心方法包括：').bold = True
for m in [
    '5-Whys（连续追问法）：通过连续5次追问"为什么"深入根因',
    '5W2H（七问分析法）：What/When/Where/Why/Who/How/How Much全面分析',
    'E-C失效机理分析（Effect-Cause）：从激活事件到故障症状的正向"技术回放"',
    'Fishbone/Ishikawa（鱼骨图）：按人机料法环分类归因',
    'FMEA（失效模式与影响分析）：系统化识别潜在失效模式',
]:
    doc.add_paragraph(m, style='List Bullet')

p = doc.add_paragraph()
p.add_run('RCA四阶段：').bold = True
p.add_run('收集信息 → 理解问题 → 确定根本原因 → 制定解决方案')
p = doc.add_paragraph()
p.add_run('成熟度：🟢 已规模化').bold = True

doc.add_heading('2.1.3 缺陷趋势分析法', level=3)
doc.add_paragraph('除ODC和RCA外，学术界还发展了多种缺陷趋势分析方法：[58][60][61]')

add_table(doc,
    ['方法', '核心思想', '适用场景', '成熟度'],
    [
        ['Gompertz模型', '基于Gompertz曲线拟合缺陷增长趋势，评估测试充分性', '测试退出判断', '🟢已规模化'],
        ['Rayleigh模型', '通过生命周期各阶段缺陷发现率建模，预测软件质量', '质量预测', '🟢已规模化'],
        ['四象限分析法', '按累积时间与缺陷去除率划分子系统质量区间', '测试策略调整', '🟡试点中'],
        ['FST缺陷流出分析', '按缺陷引入与流出维度反向推动流程改进', '过程改进', '🟡试点中'],
    ],
    col_widths=[3, 5, 3, 2.5]
)

doc.add_heading('2.1.4 软件缺陷分类标准体系', level=3)
doc.add_paragraph('业界存在多个软件缺陷分类标准，华为云Uncle_Tom在《细数应用软件的缺陷分类》[62]中系统梳理了主要标准：')

add_table(doc,
    ['标准', '发布组织', '核心内容', '分类数', '适用性'],
    [
        ['GB/T 30279-2020', '国标委', '网络安全漏洞分类分级指南', '26类', '安全漏洞分类'],
        ['CWE-1003', 'MITRE', '弱点映射到已发布漏洞的简化视图', '130类/37一级', '漏洞分类'],
        ['CWE-699', 'MITRE', '开发者视图（软件生命周期角度）', '399类/40分类', '开发缺陷分类'],
        ['CWE-1000', 'MITRE', '研究者视图（软件行为抽象）', '933类/10分类', '学术研究'],
        ['CWE-1400', 'MITRE', '软件安全保障分类（互斥分类）', '建议参考', '工具评估'],
        ['IEEE 1044-2009', 'IEEE', '软件异常分类标准', '多维度', '通用软件异常'],
        ['NVD分类', 'NIST', '参考CWE对CVE进行分类', '随CWE更新', '漏洞数据库'],
    ],
    col_widths=[2.5, 2, 4, 2, 2.5]
)

doc.add_heading('2.2 最新研究进展', level=2)

doc.add_heading('2.2.1 基于机器学习的自动化缺陷分类', level=3)
p = doc.add_paragraph()
p.add_run('核心思路：').bold = True
p.add_run('利用NLP和深度学习技术，自动对缺陷报告进行分类、标签化和优先级排序')
p = doc.add_paragraph()
p.add_run('代表论文：').bold = True
for paper in [
    'Deep learning-based software bug classification [4] — 使用深度学习自动分类软件Bug',
    'Multilabel classification for defect prediction in software engineering [5] — 多标签缺陷预测',
    'Deeplinedp: Towards a deep learning approach for line-level defect prediction [6] — 行级缺陷预测',
    'Bug priority prediction technique based on intuitionistic fuzzy representation [8] — 缺陷优先级预测',
]:
    doc.add_paragraph(paper, style='List Bullet')
p = doc.add_paragraph()
p.add_run('成熟度：🟡试点中（学术验证阶段，部分大厂开始内部试点）').bold = True

doc.add_heading('2.2.2 AIOps根因分析', level=3)
p = doc.add_paragraph()
p.add_run('核心思路：').bold = True
p.add_run('利用图算法、聚类算法和机器学习实现告警/故障的自动化根因定位')
p = doc.add_paragraph()
p.add_run('代表论文/框架：').bold = True
for paper in [
    'HRCA: A heterogeneous graph-based adaptive root cause analysis framework [11] — 异构图自适应根因分析',
    'Clustering Intrusion Detection Alarms to Support Root Cause Analysis (Julisch, 2002) [2] — 告警聚类根因分析',
    'Fault diagnosis for test alarms in microservices through multi-source data [9] — 华为微服务故障诊断',
    'What causes my test alarm? Automatic cause analysis [10] — 华为集成测试告警根因分析',
]:
    doc.add_paragraph(paper, style='List Bullet')
p = doc.add_paragraph()
p.add_run('成熟度：🟡试点中（部分大厂如美团、去哪儿已落地，但通用性待验证）').bold = True

doc.add_heading('2.2.3 数据库缺陷实证研究', level=3)
p = doc.add_paragraph()
p.add_run('核心发现：').bold = True
p.add_run('OceanBase联合中国人民大学对MySQL、SQLite、openGauss三大数据库的777个缺陷进行了实证研究，构建了四维分析框架：[47]')
for dim in [
    '根因维度：12类（错误逻辑32.3%、类型处理9.0%、API误用8.4%等）',
    '症状维度：错误结果42.99%、崩溃、死锁、性能退化等',
    '模块维度：解析器、优化器、执行引擎、存储层等',
    '关联性：类型相关根因多导致错误结果，集中于表达式求值模块',
]:
    doc.add_paragraph(dim, style='List Bullet')
p = doc.add_paragraph()
p.add_run('成熟度：🥈近三年高引用（IEEE TSE录用）').bold = True

doc.add_heading('2.2.4 大模型辅助缺陷检测与归因', level=3)
p = doc.add_paragraph()
p.add_run('核心思路：').bold = True
p.add_run('利用LLM（大语言模型）辅助缺陷检测、分类和归因分析')
p = doc.add_paragraph()
p.add_run('代表研究：').bold = True
for paper in [
    'TestGPT-Server: Automatically Testing Microservices with LLMs at ByteDance [13] — 字节跳动LLM自动测试',
    'Bitsai-cr: Automated code review via LLM in practice [14] — 字节LLM代码审查实践',
    'Effort-aware JIT defect identification at Alibaba [3] — 阿里巴巴即时缺陷识别',
    '大模型在代码缺陷检测领域的应用实践（百度Geek说）[56] — 百度大模型缺陷检测实践',
]:
    doc.add_paragraph(paper, style='List Bullet')
p = doc.add_paragraph()
p.add_run('成熟度：🔴探索中（学术和早期试点阶段）').bold = True

doc.add_heading('2.3 学术界的共识与争议', level=2)
p = doc.add_paragraph()
p.add_run('共识：').bold = True
for c in [
    '多维分类优于单维分类：ODC的多维正交分类法已被广泛验证，单一维度分类无法满足过程改进需求',
    '根因分析必须结构化：5-Whys、鱼骨图、RCA等方法已成为行业标准实践',
    'AI辅助分类是明确趋势：所有最新研究都指向AI/ML在缺陷自动分类和归因中的应用',
]:
    doc.add_paragraph(c, style='List Number')

p = doc.add_paragraph()
p.add_run('争议：').bold = True
for d in [
    '分类维度数量选择：ODC 8维度114类太细vs简化版3-5维度实用主义的争论',
    'AI自动分类的可解释性：深度学习分类结果缺乏可解释性，影响工程师信任度',
    '统一标准vs领域适配：是否应该建立行业统一标准，还是允许各领域定制化',
]:
    doc.add_paragraph(d, style='List Number')

doc.add_page_break()

# Save part 1
outpath = os.path.join(OUTDIR, 'report_part1.docx')
doc.save(outpath)
print(f'Part 1 saved: {outpath}')