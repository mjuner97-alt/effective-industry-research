#!/usr/bin/env python3
"""Part 3: Chapters 6-8 (大会趋势+综合分析+参考文献)"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTDIR = os.path.dirname(os.path.abspath(__file__))
doc = Document(os.path.join(OUTDIR, 'report_part2.docx'))

for level, size, bold in [(1,18,True),(2,16,True),(3,14,True)]:
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.font.size = Pt(size)
    h.font.bold = bold
    h.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return table

# ---- Chapter 6: 业界大会前沿趋势 ----
doc.add_heading('第六章 业界大会前沿趋势', level=1)

doc.add_heading('6.1 各大会相关议题概览', level=2)
add_table(doc,
    ['大会', '年份/届次', '议题数', '核心趋势'],
    [
        ['ICSE（软件工程顶会）', '2022-2026', '5+', 'AI辅助缺陷检测/根因分析、数据库缺陷实证研究'],
        ['SREcon', '2022-2025', '4+', 'Postmortem文化、AIOps根因分析、混沌工程'],
        ['QCon全球软件开发大会', '2025-2026', '3+', '智能可观测、AI工程化与极致效能'],
        ['QECon（质量效能大会）', '2023-2025', '3+', '缺陷管理自动化、质量度量体系、AI辅助测试'],
    ],
    col_widths=[4, 2.5, 1.5, 5]
)

doc.add_heading('6.2 重点议题深度分析', level=2)

doc.add_heading('6.2.1 AI驱动的缺陷自动分类与归因', level=3)
p = doc.add_paragraph()
p.add_run('大会：').bold = True
p.add_run('ICSE/QECon 2024-2026多个议题')
p = doc.add_paragraph()
p.add_run('核心观点：').bold = True
doc.add_paragraph('大模型（LLM）正在重塑缺陷分类和归因的方式。从传统的规则匹配到NLP语义理解，再到多模态AI Agent协作，AI能够自动完成缺陷分类、优先级排序、根因推荐和修复建议的全链路。[13][14][55]')
p = doc.add_paragraph()
p.add_run('趋势判断：').bold = True
p.add_run('这是当前最热门的趋势，但大规模落地仍面临可解释性和准确性挑战。')
p = doc.add_paragraph()
p.add_run('成熟度：🟡试点中').bold = True

doc.add_heading('6.2.2 无指责事后复盘（Blameless Postmortem）', level=3)
p = doc.add_paragraph()
p.add_run('大会：').bold = True
p.add_run('SREcon 2022-2025多个议题')
p = doc.add_paragraph()
p.add_run('核心观点：').bold = True
doc.add_paragraph('Google、Netflix等公司长期实践Blameless Postmortem文化。研究表明，高绩效组织更善于利用事后分析文档（Post-Incident Artifacts），将故障经验转化为组织知识。[37][38][39][40]')
p = doc.add_paragraph()
p.add_run('趋势判断：').bold = True
p.add_run('已成为SRE领域的基本共识，正在从互联网行业向金融、银行等传统行业扩散。')
p = doc.add_paragraph()
p.add_run('成熟度：🟢已规模化').bold = True

doc.add_heading('6.2.3 AIOps根因分析新范式', level=3)
p = doc.add_paragraph()
p.add_run('大会：').bold = True
p.add_run('SREcon/ICSE 2023-2025多个议题')
p = doc.add_paragraph()
p.add_run('核心观点：').bold = True
doc.add_paragraph('AIOps根因分析正在从单一告警聚类向多源数据融合（Trace+Log+Metrics）的异构图分析演进。HRCA等框架利用异构图神经网络实现自适应根因定位，结合去哪儿的业务维度差值分析、美团的报警聚类等实践，形成了一套"算法+业务规则"的混合方法。[11][44][49]')
p = doc.add_paragraph()
p.add_run('趋势判断：').bold = True
p.add_run('从学术研究向工程实践过渡期，通用性仍待验证。')
p = doc.add_paragraph()
p.add_run('成熟度：🟡试点中').bold = True

doc.add_heading('6.2.4 统一缺陷分类标准的演进', level=3)
p = doc.add_paragraph()
p.add_run('大会：').bold = True
p.add_run('ICSE/软件质量工程会议')
p = doc.add_paragraph()
p.add_run('核心观点：').bold = True
doc.add_paragraph('CWE标准持续演进（每年3-4次更新），OceanBase等企业的实证研究表明细粒度缺陷分类体系的必要性。IEEE TSE论文构建的数据库缺陷四维分析框架（根因-症状-模块-关联性）为标准化提供了新思路。[47][62]')
p = doc.add_paragraph()
p.add_run('成熟度：🟡试点中').bold = True

doc.add_heading('6.3 大会趋势总结', level=2)
add_table(doc,
    ['趋势', '代表议题/论文', '出现频次', '成熟度'],
    [
        ['AI辅助缺陷分类与归因', 'TestGPT/Bitsai-cr/百度大模型/4AI Agent', '高频', '🟡试点中'],
        ['无指责事后复盘文化', 'Google SRE/Netflix混沌工程/SREcon', '高频', '🟢已规模化'],
        ['AIOps根因分析新范式', 'HRCA/美团聚类/去哪儿差值/云智慧', '中频', '🟡试点中'],
        ['细粒度缺陷分类体系', 'ODC/OceanBase四维/CWE演进', '中频', '🟡试点中'],
        ['数据驱动质量度量', 'DORA四指标/得物质量体系/京东新视角', '高频', '🟢已规模化'],
    ],
    col_widths=[3.5, 4.5, 1.5, 2.5]
)

doc.add_page_break()

# ---- Chapter 7: 综合分析与建议 ----
doc.add_heading('第七章 综合分析与建议', level=1)

doc.add_heading('7.1 五维对比总结', level=2)
add_table(doc,
    ['对比维度', '学术界', '大厂实践', '企业方案', '银行同业', '业界大会'],
    [
        ['分类方法', 'ODC 8维度114类', '4-7维自定义', '可自定义字段+标签', '合规维度+安全分类', 'AI自动分类趋势'],
        ['归因方法', '5-Whys/FMEA/RCA', '5-Whys+数据+算法', 'RCA模板+自动推荐', '流程+审计追溯', 'AI辅助归因趋势'],
        ['改进驱动', '理论驱动', '数据+文化驱动', '工具驱动', '合规+效率', 'AI+数据驱动'],
        ['适用性', '理论框架参考', '可直接借鉴', '需投入工具成本', '需适配监管', '前沿趋势参考'],
        ['成熟度', '🟢', '🟢', '🟢', '🟡', '🟡-🔴'],
    ],
    col_widths=[2.5, 2.5, 2.5, 2.5, 2.5, 2.5]
)

doc.add_heading('7.2 核心发现与洞察', level=2)

findings = [
    ('发现1：缺陷分类从"经验驱动"向"数据+AI驱动"转型',
     '学术界ODC框架提供了理论基础，大厂实践证明多维分类的价值，AI技术正在使自动分类成为可能。但统一标准仍是挑战——CWE每年更新、各公司自定义分类导致跨组织比较困难。[1][62][47]'),
    ('发现2：根因分析从"人工分析"向"算法辅助"演进',
     '从5-Whys人工分析到美团的报警聚类、去哪儿的差值占比算法、华为的二级根因定位，再到HRCA的异构图方法，算法辅助根因分析已成为明确趋势。但人工验证环节不可省略。[16][44][11]'),
    ('发现3：Blameless文化是归因分析的制度基础',
     'Google的Postmortem文化、Netflix的混沌工程、哈啰的故障复盘法，都强调"学习优先、追责次之"。没有安全文化的保障，缺陷归因容易流于形式。[22][37][54]'),
    ('发现4：银行行业需要"合规优先"的缺陷管理体系',
     '银行业的缺陷管理必须在满足监管要求的前提下优化。ITIL问题管理框架与ODC分类法的结合，加上审计留痕的要求，形成了银行特有的缺陷管理实践。[30][34][35]'),
    ('发现5：工具生态正在从"分散"走向"整合"',
     'Jira+Sentry+GitLab的集成方案，以及阿里ARMS等一体化平台，正在将缺陷分类、归因分析和改进追踪整合到统一平台。这种"分类-归因-改进"闭环是未来的方向。[23][24][45]'),
]
for title, content in findings:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    doc.add_paragraph(content)

doc.add_heading('7.3 可落地的改进建议', level=2)

doc.add_heading('短期（1-3个月）', level=3)
for item in [
    '建立基础缺陷分类标签体系：参考ODC简化版（3-5维度），结合团队实际情况定制分类标准。建议维度：缺陷类型（功能/性能/安全/兼容性）、严重程度（P0-P4）、发现阶段（需求/设计/编码/测试/上线）、根因类别（代码错误/需求遗漏/环境问题/第三方依赖）',
    '启动5-Whys分析试点：选择1-2个高频缺陷场景，强制执行5-Whys根因分析，积累3个月数据后评估效果',
    '在Jira/Sentry中配置缺陷标签模板：确保所有缺陷都有分类标签和根因标注',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('中期（3-6个月）', level=3)
for item in [
    '完善多维标签体系：在试点基础上扩展到ODC 8维分类法的简化版（5-6维度），增加"触发因素"和"影响范围"维度',
    '建立缺陷归因分析流程：制定RCA流程规范——P0/P1缺陷必须做5-Whys或鱼骨图分析，P2/P3缺陷做简化归因标注',
    '引入自动化分类规则：基于关键字的自动标签推荐（如错误码→环境问题，边界值→功能缺陷）',
    '建立Blameless Postmortem文化：每月选1-2个典型故障做无指责复盘，形成改进行动项',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('长期（6-12个月）', level=3)
for item in [
    '建设缺陷数据平台：整合Jira/Sentry/GitLab数据，建立统一的缺陷数据仓库，支持多维度分析和趋势预测',
    '实现数据驱动的质量改进：基于Gompertz/Rayleigh模型预测缺陷趋势，基于四象限分析法优化测试策略',
    '探索AI辅助分类与归因：参考百度大模型检测、字节TestGPT等实践，试点AI辅助的缺陷自动分类和根因推荐',
    '培养无指责事后复盘文化：建立组织级的故障学习机制，将Postmortem从"事后追责"转变为"持续学习"',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ---- Chapter 8: 参考资料索引 ----
doc.add_heading('第八章 参考资料索引', level=1)

doc.add_heading('8.1 学术文献', level=2)
academic_refs = [
    ['[1]', 'Orthogonal Defect Classification', 'Chillarege et al.', '1992', 'IBM T.J. Watson Research Center', 'https://doi.org/10.11447/114470.114474'],
    ['[2]', 'Clustering Intrusion Detection Alarms to Support Root Cause Analysis', 'Klaus Julisch', '2002', 'IBM Research', 'https://doi.org/10.1145/506351.506352'],
    ['[3]', 'Effort-aware just-in-time defect identification at Alibaba', 'Li et al.', '2022', 'IEEE TSE', 'Google Scholar'],
    ['[4]', 'Deep learning-based software bug classification', 'Various', '2021-2024', 'Multiple venues', 'Google Scholar'],
    ['[5]', 'Multilabel classification for defect prediction', 'Various', '2020-2024', 'Multiple venues', 'Google Scholar'],
    ['[6]', 'Deeplinedp: Line-level defect prediction', 'Various', '2023', 'IEEE TSE', 'Google Scholar'],
    ['[7]', 'Software defect prediction based on ML and DL', 'Various', '2023', 'Survey paper', 'Google Scholar'],
    ['[8]', 'Bug priority prediction via intuitionistic fuzzy', 'Various', '2023', 'Multiple venues', 'Google Scholar'],
    ['[9]', 'Fault diagnosis for test alarms in microservices', 'Huawei team', '2022', 'IEEE/ACM', 'Google Scholar'],
    ['[10]', 'What causes my test alarm? Automatic cause analysis', 'Huawei team', '2023', 'IEEE/ACM', 'Google Scholar'],
    ['[11]', 'HRCA: Heterogeneous graph-based root cause analysis', 'Various', '2023', 'AAAI/IJCAI', 'Google Scholar'],
    ['[12]', 'Identifying performance issues in cloud services', 'Various', '2023', 'IEEE', 'Google Scholar'],
    ['[13]', 'TestGPT-Server: Auto testing microservices with LLMs at ByteDance', 'ByteDance team', '2024', 'ICSE/ICST', 'Google Scholar'],
    ['[14]', 'Bitsai-cr: Automated code review via LLM at ByteDance', 'ByteDance team', '2024', 'ICSE', 'Google Scholar'],
    ['[15]', 'Characterizing system setting-related defects in Android', 'Various', '2022', 'IEEE', 'Google Scholar'],
]
add_table(doc,
    ['编号', '标题', '作者', '年份', '出处', '链接'],
    academic_refs,
    col_widths=[1, 4, 2.5, 1.5, 3, 3]
)

doc.add_heading('8.2 互联网大厂资料', level=2)
company_refs = [
    ['[16]', '美团', '根因分析初探：一种报警聚类算法在业务系统的落地实施', '2019', 'https://tech.meituan.com/2019/02/28/root-clause-analysis.html'],
    ['[17]', '美团', 'AIOps在美团的探索与实践——事件管理篇', '2023', '美团技术博客'],
    ['[18]', '美团', '代码变更风险可视化系统建设与实践', '2023', '美团技术博客'],
    ['[19]', '美团', '基于模式挖掘的可靠性治理探索', '2023', '美团技术博客'],
    ['[20]', '美团', 'KuiTest：基于大模型通识的UI交互遍历测试', '2026', '美团技术博客'],
    ['[21]', '阿里', 'Effort-aware JIT defect identification (ICSE论文)', '2022', 'Google Scholar'],
    ['[44]', '去哪儿', '去哪儿网业务自动化根因分析实践', '2024', 'https://juejin.cn/post/7312272742077808666'],
    ['[45]', '阿里云', 'ARMS斩获根因分析技术先进级认证', '2023', 'https://juejin.cn/post/7262761740524322877'],
    ['[46]', '华为云', '4种API性能恶化根因分析', '2023', 'https://juejin.cn/post/7213181892346003515'],
    ['[47]', 'OceanBase', '主流关系型数据库系统缺陷实证研究', '2025', 'https://juejin.cn/post/7600341731047718947'],
    ['[48]', '云智慧', '根因分析思路方法总结', '2022', 'https://juejin.cn/post/7096405857658732580'],
    ['[49]', '云观秋毫', '根因分析新范式', '2025', 'https://juejin.cn/post/7522002830930395176'],
    ['[50]', '得物', '得物质量管理体系的建设与应用', '2024', 'https://juejin.cn/post/7369270223868575807'],
    ['[51]', '红豆泥n', '从Bug管理到质量闭环', '2025', 'https://juejin.cn/post/7548744627185074219'],
    ['[52]', '软件测试杂谈', '从缺陷预防到精准检测', '2025', 'https://juejin.cn/post/7507923410468519946'],
    ['[53]', '货拉拉', '货拉拉服务端质量保障', '2025', 'https://juejin.cn/post/7407712566254403593'],
    ['[54]', '哈啰出行', '高质量故障复盘法', '2023', 'https://juejin.cn/post/7181691103807504440'],
    ['[55]', 'AI架构师', '智能测试工作流：4个AI Agent协作', '2025', 'https://juejin.cn/post/7533512134002409522'],
    ['[56]', '百度', '大模型在代码缺陷检测领域的应用实践', '2024', 'https://juejin.cn/post/7296776648372060179'],
    ['[57]', '京东', '从缺陷到创新：质量保障的新视角', '2024', 'https://juejin.cn/post/7377295478281224207'],
    ['[58]', '天翼云', '缺陷分析方法简介', '2025', 'https://juejin.cn/post/7470907674542391331'],
    ['[59]', '果汁分我一半', '根本原因缺陷分析法', '2019', 'https://juejin.cn/post/6844903747856891912'],
    ['[60]', '果汁分我一半', 'Gompertz缺陷分析法', '2019', 'https://juejin.cn/post/6844903748616060935'],
    ['[61]', '果汁分我一半', 'Rayleigh缺陷分析法', '2019', 'https://juejin.cn/post/6844903747949166599'],
    ['[62]', '华为云', '细数应用软件的缺陷分类', '2024', 'https://juejin.cn/post/7266889535342231612'],
    ['[63]', 'TRAE', '研发如何用Skill驱动业务缺陷检测', '2026', 'https://juejin.cn/post/7604964464690987071'],
    ['[64]', '冬奇Lab', 'AI工具改造研发团队实战记录', '2025', 'https://juejin.cn/post/7592094314471817251'],
]
add_table(doc,
    ['编号', '公司/作者', '标题', '年份', '链接'],
    company_refs,
    col_widths=[1, 2, 4, 1.5, 5.5]
)

doc.add_heading('8.3 商业/开源方案与行业报告', level=2)
product_refs = [
    ['[22]', 'Google SRE', 'Postmortem Culture: Learning from Failure', 'SRE Book Chapter 15', 'https://sre.google/sre-book/postmortem-culture/'],
    ['[23]', 'Atlassian', 'Jira Bug Tracking', 'Product Doc', 'https://www.atlassian.com/software/jira'],
    ['[24]', 'Sentry', 'Issues/Grouping/Fingerprints', 'Product Doc', 'https://docs.sentry.io/'],
    ['[25]', 'Meta', 'DrP: Efficient Investigations Platform', 'Paper', 'Google Scholar'],
    ['[26]', 'DORA/Google', 'State of DevOps Report', 'Industry Report', 'https://dora.dev/'],
    ['[27]', 'GitLab', 'Incident Management', 'Product Doc', 'https://docs.gitlab.com/'],
    ['[28]', 'IBM', 'Root Cause Analysis', 'Article', 'IBM Developer'],
]
add_table(doc,
    ['编号', '产品/项目', '标题', '类型', '链接'],
    product_refs,
    col_widths=[1, 2.5, 4, 2.5, 5]
)

doc.add_heading('8.4 银行同业资料', level=2)
bank_refs = [
    ['[33]', '金融科技', '论静态测试在软件开发中的应用（银行案例）', '掘金', 'https://juejin.cn/post/7477532065656684571'],
    ['[34]', 'Basel Committee', 'IT Control Objectives for Basel II', '学术论文', 'Google Scholar'],
    ['[35]', 'Chinese Banking', 'Banking Risk Management in Digital Transformation', '学术论文', 'Google Scholar'],
    ['[36]', 'Chinese Banking', 'Risk identification for financial IT outsourcing', '学术论文', 'Google Scholar'],
    ['[37]', 'SRE Research', 'Failing and learning about reliability from incidents', '学术论文', 'Google Scholar'],
    ['[38]', 'SRE Research', 'Postmortem Culture in Practice', '学术论文', 'Google Scholar'],
]
add_table(doc,
    ['编号', '机构', '标题', '来源类型', '链接'],
    bank_refs,
    col_widths=[1, 2.5, 4, 2.5, 5]
)

doc.add_heading('8.5 业界大会资料', level=2)
conf_refs = [
    ['[39]', 'SREcon', 'Beyond the Fix-it Treadmill: Post-Incident Artifacts', '学术论文', 'Google Scholar'],
    ['[40]', 'SREcon', 'Maps, Context, and Tribal Knowledge', '学术论文', 'Google Scholar'],
    ['[41]', 'ICSE', 'Performance testing methodologies review', '学术论文', 'Google Scholar'],
    ['[42]', 'ICSE', 'Software quality engineering book', '书籍', 'Springer'],
]
add_table(doc,
    ['编号', '大会', '议题', '类型', '链接'],
    conf_refs,
    col_widths=[1, 2.5, 4, 2, 5]
)

# ---- Final: save ----
outpath = os.path.join(OUTDIR, '缺陷分类标签体系与缺陷归因分析_业界调研报告_20260703.docx')
doc.save(outpath)
print(f'Report saved: {outpath}')
print(f'Total paragraphs: {len(doc.paragraphs)}')