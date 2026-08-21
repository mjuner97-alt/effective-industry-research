#!/usr/bin/env python3
"""Part 2: Chapters 3-5 (大厂实践+国外企业+银行同业)"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTDIR = os.path.dirname(os.path.abspath(__file__))
doc = Document(os.path.join(OUTDIR, 'report_part1.docx'))

# Fix styles
for level, size, bold in [(1,18,True),(2,16,True),(3,14,True)]:
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.font.size = Pt(size)
    h.font.bold = bold
    h.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return table

# ---- Chapter 3: 大厂实践 ----
doc.add_heading('第三章 互联网大厂实践', level=1)

doc.add_heading('3.1 国内互联网大厂做法', level=2)

doc.add_heading('3.1.1 美团', level=3)
p = doc.add_paragraph()
p.add_run('缺陷管理流程：').bold = True
p.add_run('美团建立了从缺陷发现到闭环的完整流程，涵盖报警收敛、聚类分析、根因定位和改进追踪。[16]')
p = doc.add_paragraph()
p.add_run('分类方法：').bold = True
p.add_run('采用基于报警聚类的根因分析方法，提取5个关键特征维度：机房、环境、异常来源、报警文本关键内容、故障位置（接口/类）。使用泛化层次结构（DAG）描述属性泛化关系，min_size=1/5×报警数量，ε=0.05。[16]')
p = doc.add_paragraph()
p.add_run('归因分析方法：').bold = True
p.add_run('基于Julisch 2002论文的报警聚类算法，将具有相同根因的报警归为泛化报警，实现从数百条报警快速收敛至3-5个根因。后续升级为AIOps事件管理系统，支持事件关联、根因推荐和自动处置。[16][17]')
p = doc.add_paragraph()
p.add_run('质量度量体系：').bold = True
p.add_run('报警收敛率、根因定位准确率（>95%）、MTTR（平均恢复时间）')
p = doc.add_paragraph()
p.add_run('可借鉴点：').bold = True
for item in ['报警聚类+泛化层次结构的根因分析方法可推广', '从人工分析到自动化分析的演进路径清晰', 'AIOps事件管理系统的架构设计可参考']:
    doc.add_paragraph(item, style='List Bullet')
p = doc.add_paragraph()
p.add_run('成熟度：🟢已规模化').bold = True
p = doc.add_paragraph()
p.add_run('来源：[16][17]').font.size = Pt(9)

doc.add_heading('3.1.2 去哪儿', level=3)
p = doc.add_paragraph()
p.add_run('归因分析方法：').bold = True
p.add_run('去哪儿在交易拦截可视化项目中实现了自动化根因分析（RCA），核心思路是将根因数据抽象为"业务属性+系统错误码"的组合维度，通过比较故障前后维度差值占比定位根因。[44]')
p = doc.add_paragraph()
p.add_run('核心算法：').bold = True
p.add_run('参考Adtributor算法但做了简化——放弃维度波动性（惊奇度），只用差值占比表达维度相关性。三个简化假设：① 未来时间点总量不变；② 一般是一种问题而非一系列根因；③ 增量部分占比大。')
p = doc.add_paragraph()
p.add_run('实施效果：').bold = True
p.add_run('秒级根因分析，准确率95%以上；新接入场景2小时配置即可具备完整根因分析能力；全年预计人效提升1500人天。[44]')
p = doc.add_paragraph()
p.add_run('可借鉴点：').bold = True
for item in ['维度差值占比的简化根因算法，易于理解和实施', '业务维度+错误码的根因数据设计思路', '小时级接入新场景的通用化方案']:
    doc.add_paragraph(item, style='List Bullet')
p = doc.add_paragraph()
p.add_run('成熟度：🟢已规模化').bold = True
p = doc.add_paragraph()
p.add_run('来源：[44]').font.size = Pt(9)

doc.add_heading('3.1.3 阿里/蚂蚁', level=3)
p = doc.add_paragraph()
p.add_run('分类方法：').bold = True
p.add_run('阿里巴巴在即时缺陷预测（JIT defect identification）方面有深入研究，通过代码变更特征和开发者上下文信息实现缺陷预测。[3][21]')
p = doc.add_paragraph()
p.add_run('归因分析：').bold = True
p.add_run('阿里云ARMS（应用实时监控服务）在2023年通过了中国信通院"根因分析技术分级能力要求"先进级认证，是业界最早获得此类认证的产品之一。[45]')
p = doc.add_paragraph()
p.add_run('成熟度：🟢已规模化').bold = True
p = doc.add_paragraph()
p.add_run('来源：[3][21][45]').font.size = Pt(9)

doc.add_heading('3.1.4 华为', level=3)
p = doc.add_paragraph()
p.add_run('分类方法：').bold = True
p.add_run('华为云Uncle_Tom系统梳理了GB/T 30279、CNNVD、NVD、CWE等缺陷分类标准，为团队选择合适的缺陷分类方法提供了参考框架。[62]')
p = doc.add_paragraph()
p.add_run('归因分析方法：').bold = True
p.add_run('华为提出二级根因分析方法：先在异常调用链内部分析候选根因，再在全局拓扑环境下对候选根因进行汇聚，实现API性能恶化的快速根因定位。[46] 此外，华为在ICSE等顶会发表多篇根因分析论文，涉及微服务测试告警故障诊断[9]和集成测试告警根因分析[10]。')
p = doc.add_paragraph()
p.add_run('产品：').bold = True
p.add_run('华为云CodeArts Defect缺陷管理服务，提供结构化缺陷跟踪流程和标准化的质量度量模型。[23]')
p = doc.add_paragraph()
p.add_run('成熟度：🟢已规模化').bold = True
p = doc.add_paragraph()
p.add_run('来源：[9][10][46][62][23]').font.size = Pt(9)

doc.add_heading('3.1.5 百度/京东/得物/货拉拉/哈啰', level=3)
add_table(doc,
    ['公司', '核心实践', '关键特点', '成熟度', '来源'],
    [
        ['百度', '大模型在代码缺陷检测领域的应用实践', '静态代码扫描(SA)+大模型自动生成检测规则', '🟡试点中', '[56]'],
        ['京东', '从缺陷到创新：质量保障的新视角', '质量保障体系化建设', '🟢已规模化', '[57]'],
        ['得物', '质量管理体系的建设与应用', '渐进式质量保障体系，场景驱动', '🟢已规模化', '[50]'],
        ['货拉拉', '服务端质量保障之测试策略篇', '测试策略体系化', '🟢已规模化', '[53]'],
        ['哈啰出行', '高质量故障复盘法', '结构化故障复盘方法论', '🟢已规模化', '[54]'],
    ],
    col_widths=[2, 4, 4, 2.5, 1.5]
)

doc.add_heading('3.1.6 字节跳动', level=3)
p = doc.add_paragraph()
p.add_run('缺陷检测：').bold = True
p.add_run('字节跳动在AI辅助测试方面有两项重要实践：TestGPT-Server（LLM自动测试微服务）和Bitsai-cr（LLM代码审查），代表了大模型在缺陷检测领域的最新探索。[13][14]')
p = doc.add_paragraph()
p.add_run('成熟度：🔴探索中').bold = True

doc.add_heading('3.2 国外科技公司做法', level=2)

doc.add_heading('3.2.1 Google', level=3)
p = doc.add_paragraph()
p.add_run('RCA/事后复盘方法：').bold = True
p.add_run('Google的Postmortem文化是业界标杆。核心原则是"无指责复盘"（Blameless Postmortem），强调学习优先、追责次之。事后复盘文档包含：事件记录、影响范围、缓解/解决措施、根因分析、后续行动项。[22]')
p = doc.add_paragraph()
p.add_run('可借鉴点：').bold = True
for item in ['无指责复盘文化——鼓励工程师主动报告和分享', '模板化Postmortem文档确保根因分析的结构化', 'DORA四项关键指标（部署频率、变更失败率、MTTR、交付周期）驱动改进']:
    doc.add_paragraph(item, style='List Bullet')
p = doc.add_paragraph()
p.add_run('成熟度：🟢已规模化').bold = True
p = doc.add_paragraph()
p.add_run('来源：[22][26]').font.size = Pt(9)

doc.add_heading('3.2.2 Meta', level=3)
p = doc.add_paragraph()
p.add_run('DrP高效调查平台：').bold = True
p.add_run('Meta开发了DrP（Drive Root-cause and Resolution Platform）平台，支持大规模事件的快速调查和根因分析，实现调查流程的标准化和自动化。[24]')
p = doc.add_paragraph()
p.add_run('成熟度：🟢已规模化').bold = True

doc.add_heading('3.2.3 Netflix', level=3)
p = doc.add_paragraph()
p.add_run('混沌工程实践：').bold = True
p.add_run('Netflix通过混沌工程（Chaos Engineering）主动注入故障来发现系统薄弱点，配合Postmortem实现从"被动响应"到"主动发现"的根因分析演进。[37]')
p = doc.add_paragraph()
p.add_run('成熟度：🟢已规模化').bold = True

doc.add_heading('3.3 共性模式总结', level=2)
add_table(doc,
    ['模式', '美团', '去哪儿', '阿里', '华为', 'Google', '说明'],
    [
        ['分类方式', '5维报警聚类', '维度+错误码', '代码变更特征', 'CWE+自定义', 'Postmortem模板', '从简单标签到多维体系'],
        ['归因方法', '聚类算法', '差值占比', 'ARMS认证', '二级根因', 'Blameless', '算法辅助→文化驱动'],
        ['改进驱动', '数据驱动', '数据驱动', '数据+AI', '数据驱动', '文化驱动', '数据+文化双轮驱动'],
        ['自动化程度', '高', '高', '中-高', '中', '中', '大厂自动化程度普遍提升'],
    ],
    col_widths=[2.5, 2, 2, 2, 2, 2, 3]
)

doc.add_heading('3.4 差异分析', level=2)
p = doc.add_paragraph()
p.add_run('国内 vs 国外：').bold = True
p.add_run('国内大厂更注重算法驱动和自动化程度，如美团的报警聚类、去哪儿的差值占比算法；国外更注重文化层面，如Google的Blameless Postmortem和Netflix的混沌工程。两者并非对立，而是互补。')
p = doc.add_paragraph()
p.add_run('不同规模：').bold = True
p.add_run('大厂（日活亿级）倾向于自建平台和算法；中小团队更适合采用简化版ODC分类+5-Whys归因+开源工具的组合方案。')
p = doc.add_paragraph()
p.add_run('不同开发模式：').bold = True
p.add_run('瀑布模式下缺陷分类偏向后期验证阶段；敏捷/DevOps模式下需要全流程嵌入分类和归因。')

doc.add_page_break()

# ---- Chapter 4: 国外企业/行业方案 ----
doc.add_heading('第四章 国外企业/行业方案', level=1)

doc.add_heading('4.1 主流商业产品方案', level=2)

doc.add_heading('4.1.1 Atlassian Jira', level=3)
p = doc.add_paragraph()
p.add_run('核心功能：').bold = True
p.add_run('缺陷追踪全流程：Capture→Assign→Prioritize→Track→Resolve')
p = doc.add_paragraph()
p.add_run('分类体系：').bold = True
p.add_run('支持自定义工作流、优先级（P1-P5）、严重性（Blocker/Trivial）、标签体系、组件分类、影响范围。可配置缺陷类型（Bug/Task/Story/Epic）。[23]')
p = doc.add_paragraph()
p.add_run('归因支持：').bold = True
p.add_run('通过插件支持RCA模板（如EasyRCA），支持5-Whys分析模板和鱼骨图。')
p = doc.add_paragraph()
p.add_run('自动化能力：').bold = True
p.add_run('基于JQL的自动化规则、Webhook通知、与Bitbucket/GitHub/Jenkins集成。')
p = doc.add_paragraph()
p.add_run('适用规模：').bold = True
p.add_run('5-5000人团队')
p = doc.add_paragraph()
p.add_run('成熟度：🟢已规模化').bold = True
p = doc.add_paragraph()
p.add_run('来源：[23]').font.size = Pt(9)

doc.add_heading('4.1.2 Sentry', level=3)
p = doc.add_paragraph()
p.add_run('核心功能：').bold = True
p.add_run('实时错误监控、自动分组（Issue Grouping）、指纹识别（Fingerprints）、状态分流（States & Triage）')
p = doc.add_paragraph()
p.add_run('分类体系：').bold = True
p.add_run('Sentry通过堆栈指纹（Stack Trace Fingerprints）自动将相似错误归为一组，支持自定义指纹规则和Issue标签。分组算法基于异常类型、消息模板和关键帧。[24]')
p = doc.add_paragraph()
p.add_run('归因支持：').bold = True
p.add_run('自动关联代码变更（Suspect Commits），显示引入错误的commit和负责人。支持分配、优先级、环境标签。')
p = doc.add_paragraph()
p.add_run('成熟度：🟢已规模化').bold = True
p = doc.add_paragraph()
p.add_run('来源：[24]').font.size = Pt(9)

doc.add_heading('4.1.3 GitLab Incident Management', level=3)
p = doc.add_paragraph()
p.add_run('核心功能：').bold = True
p.add_run('事件管理（Incident Management）、On-call轮值、告警路由、Postmortem模板')
p = doc.add_paragraph()
p.add_run('归因支持：').bold = True
p.add_run('内置Postmortem模板，支持时间线回溯、根因标注和改进行动项追踪。[26]')
p = doc.add_paragraph()
p.add_run('成熟度：🟢已规模化').bold = True

doc.add_heading('4.2 行业报告关键发现', level=2)
p = doc.add_paragraph()
p.add_run('DORA State of DevOps Report（Google主导）：').bold = True
p.add_run('DORA报告定义了四项关键指标：部署频率、变更失败率、MTTR（平均恢复时间）、交付周期。其中变更失败率直接关联缺陷分类和归因分析的质量。高效能团队的变更失败率<5%，低效能团队>15%。[26]')

doc.add_heading('4.3 ITIL问题管理框架', level=2)
p = doc.add_paragraph()
p.add_run('ITIL框架中的问题管理（Problem Management）与缺陷归因高度相关：')
for item in [
    '问题识别：通过趋势分析从重复事件中识别问题',
    '问题分类：按影响范围和紧急程度分类',
    '根因分析：使用5-Whys、鱼骨图等结构化方法',
    '已知错误数据库：建立已知错误和解决方案的知识库',
    '问题关闭：确保根因消除后关闭问题记录',
]:
    doc.add_paragraph(item, style='List Bullet')
p = doc.add_paragraph()
p.add_run('来源：[30][31][32]').font.size = Pt(9)

doc.add_page_break()

# ---- Chapter 5: 银行同业分析 ----
doc.add_heading('第五章 银行同业分析', level=1)

doc.add_heading('5.1 银行/金融机构落地案例', level=2)
p = doc.add_paragraph()
p.add_run('银行行业在缺陷管理和质量保障方面有其特殊要求和实践：')

doc.add_heading('5.1.1 金融科技质量管理', level=3)
p = doc.add_paragraph()
p.add_run('银行业IT内控与数字化转型：').bold = True
p.add_run('银行业在数字化转型中面临特殊的IT内控要求。根据《IT Control Objectives for Basel II》研究，银行需要在满足Basel II合规要求的同时建立IT治理和风险管理框架。[34]')
p = doc.add_paragraph()
p.add_run('国有商业银行数字化转型风险管理：').bold = True
p.add_run('研究分析了国有商业银行在数字化转型中的风险管理挑战和对策，强调IT风险识别和传导模型的重要性。[35]')

doc.add_heading('5.1.2 银行行业AI测试转型', level=3)
p = doc.add_paragraph()
p.add_run('掘金文章《银行行业AI测试转型实操手册》指出：银行业软件测试作为保障业务系统稳定运行的关键环节，正面临AI驱动的转型。传统测试方法的局限性逐渐凸显，AI技术为缺陷分类和归因分析提供了新思路。[32]')

doc.add_heading('5.1.3 金融科技静态测试', level=3)
p = doc.add_paragraph()
p.add_run('文琪小站的文章从金融科技企业视角论述了静态测试在软件开发中的应用，该公司为银行、证券及保险行业提供核心交易与风控系统研发服务，对缺陷管理有着严格的合规要求。[33]')

doc.add_heading('5.2 监管合规要求', level=2)
add_table(doc,
    ['监管文件', '发布机构', '核心要求', '对缺陷管理的影响'],
    [
        ['《商业银行信息科技风险管理指引》', '银保监会', '建立信息科技风险管理体系，包括事件管理、问题管理和变更管理', '缺陷分类必须覆盖合规维度，归因分析必须满足审计追溯要求'],
        ['《IT Control Objectives for Basel II》', '巴塞尔委员会', 'IT治理与风险管理框架，确保数据完整性和系统可靠性', '缺陷管理必须支持风险量化评估和监管报告'],
        ['《金融科技风险管理》', '央行等', '数字化转型中的风险识别与传导管控', '需要建立IT外包风险的缺陷传导模型'],
    ],
    col_widths=[3.5, 2, 4, 4.5]
)

doc.add_heading('5.3 银行同业特殊考量', level=2)
for item in [
    '合规优先：银行业必须在满足监管的前提下优化流程，缺陷分类标签体系需包含合规维度',
    '安全要求：数据安全、交易安全、客户隐私是银行缺陷管理的底线要求',
    '变更管理：严格的变更审批流程意味着缺陷修复需要经过完整的变更管理流程',
    '审计留痕：所有缺陷和变更需要可追溯，满足内审和外审要求',
    '外包管理：供应商交付质量的缺陷管控是银行特有的管理挑战',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

outpath = os.path.join(OUTDIR, 'report_part2.docx')
doc.save(outpath)
print(f'Part 2 saved: {outpath}')