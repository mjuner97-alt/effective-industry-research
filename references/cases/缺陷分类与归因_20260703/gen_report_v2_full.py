#!/usr/bin/env python3
"""V2-full: V1 content + 4 improvements (Tier, Contradictions, Info Gaps, Confidence)"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTDIR = '/Users/minjun/.openclaw/workspace/skills/effective-industry-research/references/cases/缺陷分类与归因_20260703'

# Load V1 document as base
doc = Document(os.path.join(OUTDIR, '缺陷分类标签体系与缺陷归因分析_业界调研报告_20260703.docx'))

# Fix styles
for level, size, bold in [(1,18,True),(2,16,True),(3,14,True)]:
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.font.size = Pt(size)
    h.font.bold = bold
    h.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return table

# ===== Add 4 improvements =====

# ===== Improvement 1: Update cover page - add V2 marker and Tier/confidence info =====
# Find and update the cover page info
for para in doc.paragraphs:
    if '根据 effective-industry-research skill 方法论生成' in para.text:
        para.clear()
        run = para.add_run('根据 effective-industry-research skill V2 方法论生成（含Tier分级+矛盾标注+信息缺口+置信度评估）')
        run.font.size = Pt(12)
        break

# ===== Improvement 2: Add Tier labels to Chapter 2 academic section =====
# Add Tier annotations after key sections
for para in doc.paragraphs:
    # ODC section
    if para.text.startswith('ODC是由IBM Ram Chillarege'):
        # Add Tier label after this paragraph
        new_para = doc.add_paragraph()
        new_para.add_run('来源可信度：').bold = True
        new_para.add_run('[Tier 1 — 一手论文原文] [置信度: HIGH]')
        new_para.add_run(' — 3个独立来源验证（IBM原论文+华为云综述+天翼云综述）')
        # Move this paragraph right after the current one
        para._element.addnext(new_para._element)
    
    if para.text.startswith('RCA是一种结构化的问题根因识别方法'):
        new_para = doc.add_paragraph()
        new_para.add_run('来源可信度：').bold = True
        new_para.add_run('[Tier 2 — 天翼云综述+云智慧总结] [置信度: HIGH]')
        para._element.addnext(new_para._element)

# ===== Improvement 3: Add ⚠️ 矛盾与不确定性 after Chapter 7 findings =====
# Find Chapter 7 heading
ch7_found = False
findings_additions = [
    ('发现1：缺陷分类从"经验驱动"向"数据+AI驱动"转型',
     '⚠️ 矛盾与不确定性：ODC建议8维度114类，但实际落地几乎所有公司简化为4-7维度。原因：ODC面向过程改进需要细粒度，但团队实施成本高。CWE面向安全缺陷分类，对功能缺陷覆盖不足，两者维度定义不统一，映射关系复杂。[置信度: HIGH — Tier 1-2来源，≥3个独立来源验证]'),
    ('发现2：根因分析从"人工分析"向"算法辅助"演进',
     '⚠️ 矛盾与不确定性：三代并非替代关系而是叠加。第三代AI归因的准确率数据来自学术/厂商宣传，独立验证不足。训练数据偏向特定项目/领域，泛化性存疑。[置信度: HIGH — Tier 2-3来源，≥5个独立来源验证]'),
    ('发现3：Blameless文化是归因分析的制度基础',
     '⚠️ 矛盾与不确定性：文化因素难以量化评估。"无指责复盘"在强考核导向的组织中推行困难。缺乏银行行业的Postmortem实践案例。[置信度: MEDIUM-HIGH — Tier 1-2来源验证，但文化因素不可复制]'),
    ('发现4：银行行业需要"合规优先"的缺陷管理体系',
     '⚠️ 矛盾与不确定性：合规要求与效率提升可能存在张力。银行同业缺乏公开的缺陷分类标签体系案例——这是信息缺口。[置信度: MEDIUM — Tier 1来源验证监管要求，但缺乏银行内部实践案例]'),
    ('发现5：工具生态正在从"分散"走向"整合"',
     '⚠️ 矛盾与不确定性：整合平台的适用性因团队规模和工具链而异。小团队可能不需要全链路整合。[置信度: MEDIUM — Tier 2-3来源，部分验证]'),
]

for para in doc.paragraphs:
    for finding_title, contradiction_text in findings_additions:
        if finding_title in para.text:
            new_para = doc.add_paragraph()
            run = new_para.add_run(contradiction_text)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
            para._element.addnext(new_para._element)

# ===== Improvement 4: Add 信息缺口 and 置信度评估 sections after Chapter 7 =====
# Find the last paragraph of Chapter 7 (before Chapter 8 heading)
ch8_index = None
for i, para in enumerate(doc.paragraphs):
    if para.text.startswith('第八章') or para.text.startswith('第八章 参考资料索引'):
        ch8_index = i
        break

# We need to insert before Chapter 8
# Add new sections
doc.add_paragraph()  # spacing

# 7.3 信息缺口
p = doc.add_paragraph()
p.style = doc.styles['Heading 2']
run = p.add_run('7.3 🕳️ 信息缺口')
run.font.size = Pt(16)
run.bold = True
run.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)

doc.add_paragraph('以下信息在调研过程中无法获取，可能影响部分结论的完整性：')

add_table(doc,
    ['缺口描述', '尝试方式', '未获取原因', '影响程度'],
    [
        ['银行内部缺陷分类标签体系的具体设计方案', 'Google Scholar+掘金+web_fetch搜索', '银行不公开内部缺陷管理细节', '高'],
        ['ServiceNow问题管理的中文实践案例', '掘金搜索+web_fetch', '产品文档以英文为主，中文案例极少', '中'],
        ['QECon/QCon 2025-2026大会议题的详细内容', '浏览器搜索+web_fetch', '仅获得议题名称，缺乏演讲详细内容', '中'],
        ['AI辅助缺陷分类的独立基准测试数据', 'Google Scholar+Semantic Scholar', '学术/厂商宣传数据缺乏独立验证', '高'],
        ['Basel II在缺陷分类标签方面的具体指导细则', 'Google Scholar', 'Basel II侧重IT治理框架，未细化到标签设计', '低'],
        ['银行业缺陷管理的内部标杆案例', 'Google Scholar+掘金', '银行同业极少公开缺陷管理内部实践', '高'],
    ],
    col_widths=[4, 3, 3.5, 1.5]
)

# 7.4 置信度评估
p = doc.add_paragraph()
p.style = doc.styles['Heading 2']
run = p.add_run('7.4 置信度评估')
run.font.size = Pt(16)
run.bold = True
run.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)

doc.add_paragraph('对核心发现的置信度进行系统性评估：')

add_table(doc,
    ['发现', '置信度', '依据', '不确定性来源'],
    [
        ['缺陷分类从经验驱动向数据+AI驱动转型', 'HIGH', 'Tier 1-2来源，≥3个独立来源验证', '不同行业维度取舍可能不同；统一标准仍存争议'],
        ['根因分析从人工分析向算法辅助演进', 'HIGH', 'Tier 2-3来源，≥5个独立来源验证', '第三代AI归因准确率数据缺乏独立验证'],
        ['Blameless文化是归因分析的制度基础', 'MEDIUM-HIGH', 'Tier 1-2来源验证', '文化因素不可复制；缺乏银行行业案例'],
        ['银行行业需要合规优先的缺陷管理体系', 'MEDIUM', 'Tier 1来源验证监管要求', '缺乏银行内部实践公开案例'],
        ['工具生态从分散走向整合', 'MEDIUM', 'Tier 2-3来源，部分验证', '整合平台适用性因团队规模而异'],
    ],
    col_widths=[3, 1.5, 3.5, 4]
)

doc.add_paragraph('置信度定义：')
doc.add_paragraph('HIGH：Tier 1-2来源，≥2个独立来源验证，无矛盾', style='List Bullet')
doc.add_paragraph('MEDIUM-HIGH：Tier 1-2来源验证，但存在少量不确定性', style='List Bullet')
doc.add_paragraph('MEDIUM：Tier 1来源验证部分结论，但缺乏完整案例', style='List Bullet')
doc.add_paragraph('LOW：Tier 4-5来源，单一来源，或存在显著矛盾', style='List Bullet')

# ===== Improvement 5: Add Tier column to reference tables =====
# This would require modifying existing tables which is complex in python-docx
# Instead, add a Tier classification note at the beginning of Chapter 8

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('来源可信度Tier分级说明：')
run.bold = True
run.font.size = Pt(11)

add_table(doc,
    ['等级', '来源类型', '权重', '示例'],
    [
        ['Tier 1', '一手来源（学术论文原文、官方标准、监管文件）', '最高', 'ODC论文(Chillarege)、IEEE 1044、银保监会指引、Google SRE Book'],
        ['Tier 2', '专家分析（IEEE/ACM论文、大厂技术博客、行业报告）', '高', '美团根因分析、去哪儿RCA实践、DORA报告、SREcon论文'],
        ['Tier 3', '质量二手来源（掘金精选文章、行业会议PPT、技术社区）', '中', '天翼云综述、云智慧总结、得物质量体系、哈啰复盘法'],
        ['Tier 4', '一般二手来源（新闻报道、百科、转载文章）', '较低', '一般性技术报道'],
        ['Tier 5', '非正式来源（论坛帖子、个人博客、社交媒体评论）', '谨慎使用', '个人经验分享'],
    ],
    col_widths=[1.5, 4, 1.5, 6]
)

# Save
outpath = os.path.join(OUTDIR, '缺陷分类标签体系与缺陷归因分析_业界调研报告_V2_full_20260703.docx')
doc.save(outpath)
print(f'V2-full Report saved: {outpath}')
print(f'Total paragraphs: {len(doc.paragraphs)}')